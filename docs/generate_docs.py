"""Generate yomi_documentation.pdf and yomi_documentation.docx"""

import os
OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ─────────────────────────────────────────────────────────────────────────
#  PDF via ReportLab Platypus
# ─────────────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor, white, black
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, KeepTogether, HRFlowable, ListFlowable, ListItem,
)
from reportlab.platypus.flowables import Flowable
from reportlab.lib import colors

# ── Palette ───────────────────────────────────────────────────────────────
ACCENT    = HexColor('#0A84FF')
DARK_BG   = HexColor('#0A0A1F')
NEAR_BG   = HexColor('#141422')
TEXT_MAIN = HexColor('#1A1A2E')
TEXT_MUTE = HexColor('#50506E')
BORDER    = HexColor('#DCDCF0')
CODE_BG   = HexColor('#0F0F1A')
CODE_FG   = HexColor('#C8C8E8')
INFO_BG   = HexColor('#EBF5FF')
INFO_FG   = HexColor('#003C78')
WARN_BG   = HexColor('#FFF3CD')
WARN_FG   = HexColor('#644600')
EVEN_ROW  = HexColor('#F5F5FC')
WHITE     = white

W, H = A4

# ── Styles ────────────────────────────────────────────────────────────────
SS = getSampleStyleSheet()

def mkstyle(**kw):
    base = kw.pop('base', 'Normal')
    name = kw.pop('name', 'custom')
    return ParagraphStyle(name=name, parent=SS[base], **kw)

S_BODY    = mkstyle(name='body',    fontName='Helvetica', fontSize=10,  leading=15, textColor=TEXT_MAIN, spaceBefore=2,  spaceAfter=6)
S_MUTED   = mkstyle(name='muted',   fontName='Helvetica', fontSize=9.5, leading=14, textColor=TEXT_MUTE, spaceAfter=4)
S_H1      = mkstyle(name='h1',      fontName='Helvetica-Bold', fontSize=22, leading=28, textColor=TEXT_MAIN, spaceBefore=12, spaceAfter=6)
S_H2      = mkstyle(name='h2',      fontName='Helvetica-Bold', fontSize=14, leading=20, textColor=TEXT_MAIN, spaceBefore=12, spaceAfter=4)
S_H3      = mkstyle(name='h3',      fontName='Helvetica-Bold', fontSize=11, leading=16, textColor=TEXT_MAIN, spaceBefore=8,  spaceAfter=3)
S_H4      = mkstyle(name='h4',      fontName='Helvetica-Bold', fontSize=9,  leading=14, textColor=TEXT_MUTE, spaceBefore=6,  spaceAfter=2, textTransform='uppercase')
S_CODE    = mkstyle(name='code',    fontName='Courier',         fontSize=8,  leading=12, textColor=CODE_FG,  backColor=CODE_BG, leftIndent=8, rightIndent=8, spaceBefore=2, spaceAfter=2)
S_INFO    = mkstyle(name='info',    fontName='Helvetica',       fontSize=9.5, leading=14, textColor=INFO_FG,  backColor=INFO_BG, leftIndent=8, rightIndent=8, spaceBefore=4, spaceAfter=4)
S_WARN    = mkstyle(name='warn',    fontName='Helvetica',       fontSize=9.5, leading=14, textColor=WARN_FG,  backColor=WARN_BG, leftIndent=8, rightIndent=8, spaceBefore=4, spaceAfter=4)
S_BULLET  = mkstyle(name='bullet',  fontName='Helvetica', fontSize=10, leading=14, textColor=TEXT_MAIN, leftIndent=16, bulletIndent=8, spaceAfter=3)
S_COVER_EYEBROW = mkstyle(name='eyebrow', fontName='Helvetica-Bold', fontSize=8,  leading=12, textColor=ACCENT, spaceAfter=4)
S_COVER_TITLE   = mkstyle(name='cover_t', fontName='Helvetica-Bold', fontSize=58, leading=64, textColor=WHITE,  spaceAfter=4)
S_COVER_SUB     = mkstyle(name='cover_s', fontName='Helvetica',      fontSize=14, leading=20, textColor=HexColor('#C0C0E0'), spaceAfter=8)
S_COVER_META    = mkstyle(name='cover_m', fontName='Helvetica',      fontSize=9,  leading=14, textColor=HexColor('#8080A0'))


# ── Cover page (drawn directly on canvas via onFirstPage callback) ─────────
def draw_cover(c):
    # Dark background
    c.setFillColor(DARK_BG)
    c.rect(0, 0, W, H, fill=1, stroke=0)
    # Accent gradient strip (top)
    for i in range(60):
        alpha = 1 - i / 60
        r = 10 + int(20 * alpha)
        g = 10 + int(30 * alpha)
        b = 31 + int(50 * alpha)
        c.setFillColorRGB(r/255, g/255, b/255)
        y = H - (i + 1) * (H * 0.4 / 60)
        c.rect(0, y, W, H * 0.4 / 60 + 1, fill=1, stroke=0)

    # Eyebrow
    c.setFillColor(ACCENT)
    c.setFont('Helvetica-Bold', 8)
    c.drawString(40, H - 80, 'TECHNICAL DOCUMENTATION')

    # Title
    c.setFillColor(WHITE)
    c.setFont('Helvetica-Bold', 68)
    c.drawString(38, H - 148, 'Yomi')
    # accent dot
    c.setFillColor(ACCENT)
    c.circle(38 + 178, H - 132, 6, fill=1, stroke=0)

    # Subtitle
    c.setFillColorRGB(0.78, 0.78, 0.90)
    c.setFont('Helvetica', 13)
    c.drawString(40, H - 170, 'Manga & manhwa reader for Android')
    c.drawString(40, H - 186, 'Developer and user reference')

    # Version pill
    pill_text = '  Version 1.0.0  ·  Flutter 3.x  ·  Dart 3.3  '
    c.setFillColorRGB(0.08, 0.16, 0.28)
    c.roundRect(38, H - 218, 260, 18, 9, fill=1, stroke=0)
    c.setStrokeColor(ACCENT)
    c.setLineWidth(0.5)
    c.roundRect(38, H - 218, 260, 18, 9, fill=0, stroke=1)
    c.setFillColor(ACCENT)
    c.setFont('Helvetica-Bold', 8.5)
    c.drawString(44, H - 211, pill_text)

    # Divider
    c.setStrokeColorRGB(0.20, 0.20, 0.35)
    c.setLineWidth(0.5)
    c.line(40, 70, W - 40, 70)

    # Meta row
    meta = [('Platform', 'Android API 21+'), ('UI', 'Cupertino'),
            ('State', 'Riverpod v2'), ('DB', 'Isar v3')]
    col_w = (W - 80) / len(meta)
    c.setFont('Helvetica-Bold', 8)
    c.setFillColorRGB(0.62, 0.62, 0.78)
    for i, (label, val) in enumerate(meta):
        c.drawString(40 + i * col_w, 60, label)
    c.setFont('Helvetica', 9)
    c.setFillColorRGB(0.88, 0.88, 0.95)
    for i, (label, val) in enumerate(meta):
        c.drawString(40 + i * col_w, 46, val)


# ── Helpers ───────────────────────────────────────────────────────────────
def sp(n=6):
    return Spacer(1, n)

def h1(num, title):
    return [
        sp(8),
        Paragraph(f'{num}&nbsp;&nbsp;{title}', S_H1),
        HRFlowable(width='100%', thickness=1.5, color=ACCENT, spaceAfter=6),
    ]

def h2(text):
    return [Paragraph(text, S_H2), HRFlowable(width='100%', thickness=0.5, color=BORDER, spaceAfter=4)]

def h3(text):
    return [Paragraph(text, S_H3)]

def h4(text):
    return [Paragraph(text.upper(), S_H4)]

def body(text):
    return Paragraph(text, S_BODY)

def bullets(items):
    li = []
    for item in items:
        # support **bold** markup
        li.append(ListItem(Paragraph(item, S_BULLET), leftIndent=20, bulletColor=ACCENT))
    return ListFlowable(li, bulletType='bullet', start='•', leftIndent=8)

def code_block(lines):
    result = []
    result.append(sp(4))
    for line in lines:
        result.append(Paragraph(line.replace(' ', '&nbsp;').replace('<', '&lt;').replace('>', '&gt;') or '&nbsp;', S_CODE))
    result.append(sp(6))
    return result

def info(text, warn=False):
    style = S_WARN if warn else S_INFO
    prefix = '⚠️  ' if warn else 'ℹ️  '
    return [sp(4), Paragraph(prefix + text, style), sp(4)]

def table(headers, rows, col_widths_pt):
    data = [headers] + rows
    t = Table(data, colWidths=col_widths_pt)
    style_cmds = [
        ('BACKGROUND', (0, 0), (-1, 0), ACCENT),
        ('TEXTCOLOR',  (0, 0), (-1, 0), WHITE),
        ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE',   (0, 0), (-1, 0), 9.5),
        ('FONTNAME',   (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE',   (0, 1), (-1, -1), 9),
        ('TEXTCOLOR',  (0, 1), (-1, -1), TEXT_MAIN),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [WHITE, EVEN_ROW]),
        ('LINEBELOW',  (0, 0), (-1, -1), 0.5, BORDER),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]
    t.setStyle(TableStyle(style_cmds))
    return [t, sp(8)]

def code_inline(text):
    return f'<font name="Courier" size="9" color="#0A0A2E">{text}</font>'


# ── Content builder ────────────────────────────────────────────────────────
def build_story():
    CW = W - 2*2.5*cm  # content width
    story = []

    # Cover page is drawn via onFirstPage callback; just advance to page 2
    story.append(PageBreak())

    # ── Section 1: Overview ──────────────────────────────────────────────
    story += h1('1', 'Overview')
    story += h2('1.1  What is Yomi?')
    story.append(body(
        'Yomi (読み — Japanese for "reading") is a feature-complete manga and manhwa '
        'reader for Android. It provides a polished, distraction-free reading experience built '
        'entirely with Flutter\'s Cupertino widget set, giving it the feel of a native iOS '
        'application while running on Android hardware.'
    ))
    story.append(body(
        'The application manages a personal library of saved titles, downloads chapters for offline '
        'reading, connects to pluggable content sources, and backs up your entire reading history '
        'to a local JSON file or Google Drive.'
    ))
    story += h2('1.2  Core Design Principles')
    story.append(bullets([
        '<b>Immersive, minimal UI</b> — Dark-first design with frosted glass overlays and tap-to-reveal chrome.',
        '<b>Reactive, no loading flickers</b> — Isar reactive streams + shimmer placeholders + instant Riverpod state propagation.',
        '<b>Pluggable sources</b> — MangaSource interface with zero build-time coupling.',
        '<b>Adaptive light & dark</b> — AppColorsX context extension resolves every colour at runtime.',
    ]))
    story.append(PageBreak())

    # ── Section 2: Features ──────────────────────────────────────────────
    story += h1('2', 'Features')

    story += h2('2.1  Library')
    story.append(body('The Library tab is the home screen. It shows every manga in a two-column grid.'))
    story.append(bullets([
        '<b>Real-time search</b> — frosted search bar filters titles as you type.',
        '<b>Filter sheet</b> — narrow by read status (All / Unread / Read / Downloaded) and genre.',
        '<b>Categories</b> — user-defined shelves. Each manga can belong to multiple categories.',
        '<b>Unread badge</b> — displayed on the card corner when there are unread chapters.',
        '<b>Shimmer skeleton</b> — animated placeholder grid during database load.',
    ]))

    story += h2('2.2  Reader')
    story.append(body('Full-screen page view with configurable direction, scale, and background.'))
    story += h4('Reading modes')
    story += table(
        ['Mode', 'Scroll direction', 'Use case'],
        [
            ['Left → Right', 'Horizontal, forward', 'Western comics, most manhwa'],
            ['Right → Left', 'Horizontal, reversed', 'Japanese manga (traditional)'],
            ['Vertical', 'Vertical', 'Webtoons, long-strip manhwa'],
        ],
        [CW*0.27, CW*0.33, CW*0.40]
    )
    story += h4('Chrome & navigation')
    story.append(bullets([
        'Tap anywhere to toggle the navigation chrome (toolbar + bottom scrubber).',
        'Drag or tap the scrubber track to jump instantly to any page.',
        'A 2 pt accent-colour progress line is always visible at the top of the screen.',
        'A page-count pill fades in after each page turn and disappears after 2 s.',
    ]))

    story += h2('2.3  Downloads')
    story.append(bullets([
        '<b>Queue</b> — chapters process one at a time with a live progress bar.',
        '<b>Pause & Resume</b> — any chapter can be paused; resumes from where it left off.',
        '<b>Cancel</b> — removes the entry; partial files are discarded.',
        '<b>Retry</b> — failed chapters show a retry button.',
    ]))
    story += h4('Per-chapter status indicators')
    story += table(
        ['Status', 'Meaning'],
        [
            ['Queued (clock)', 'Waiting in the download queue'],
            ['Downloading (spinner)', 'Actively being downloaded'],
            ['Paused', 'Download was manually paused'],
            ['Failed (✕)', 'Download failed — tap to retry'],
            ['Downloaded (✓)', 'Available offline (green tick)'],
        ],
        [CW*0.30, CW*0.70]
    )

    story += h2('2.4  Browse & Sources')
    story.append(body('The Browse tab lists all registered content sources.'))
    story += table(
        ['Source', 'Language', 'API'],
        [
            ['AsuraScans', 'English', 'JSON REST — api.asurascans.com'],
            ['AllManga', 'English', 'GraphQL — api.allanime.day'],
        ],
        [CW*0.28, CW*0.18, CW*0.54]
    )

    story += h2('2.5  Backup & Restore')
    story.append(bullets([
        '<b>Export</b> — creates a timestamped JSON file in app documents.',
        '<b>Restore</b> — merges backup into current library, preserving read/download status.',
        'Backup list screen with one-tap restore and confirmation dialog.',
    ]))
    story += info('Backup files are plain JSON and can be viewed or edited with any text editor.')

    story += h2('2.6  Settings')
    story += table(
        ['Section', 'Option', 'Values'],
        [
            ['Appearance', 'Theme', 'Dark / Light'],
            ['Library', 'Categories', 'Navigate to category management'],
            ['Reader', 'Direction', 'L→R / R→L / Vertical'],
            ['Reader', 'Page Scale', 'Fit Width / Fit Height / Original'],
            ['Reader', 'Background', 'Black / White / Sepia'],
            ['Downloads', 'Storage Location', 'Local / Google Drive'],
            ['Extensions', 'Check for Updates', 'Fetches latest GitHub release'],
            ['Backup & Sync', 'Export / Restore', 'JSON to app documents'],
        ],
        [CW*0.24, CW*0.33, CW*0.43]
    )
    story.append(PageBreak())

    # ── Section 3: Architecture ──────────────────────────────────────────
    story += h1('3', 'Architecture')

    story += h2('3.1  Tech Stack')
    story += table(
        ['Layer', 'Technology', 'Version'],
        [
            ['UI framework', 'Flutter — Cupertino widgets', '≥ 3.3'],
            ['State management', 'Riverpod', '^2.5'],
            ['Local database', 'Isar', '^3.1'],
            ['Networking', 'Dio', '^5.4'],
            ['Image loading', 'extended_image', '^8.3'],
            ['HTML parsing', 'html package', '^0.15'],
            ['Backup / archive', 'archive', '^3.6'],
            ['Google Drive', 'googleapis + google_sign_in', '^12 / ^6.2'],
            ['Platform bridge', 'MethodChannel yomi/platform', '—'],
        ],
        [CW*0.32, CW*0.50, CW*0.18]
    )

    story += h2('3.2  Project Structure')
    story += code_block([
        'lib/',
        '├── core/',
        '│   ├── database/          # Isar schemas: MangaEntry, ChapterEntry, DownloadEntry',
        '│   ├── extensions/sources/ # MangaSource plugins (AsuraScans, AllManga, ...)',
        '│   ├── providers/          # Riverpod providers',
        '│   ├── services/           # BackupService, UpdateService',
        '│   └── theme/              # AppColors, AppColorsX extension',
        '└── features/',
        '    ├── library/            # LibraryScreen, filters, category chips',
        '    ├── browse/             # BrowseScreen, SourceMangaScreen',
        '    ├── title_detail/       # TitleDetailScreen, ChapterListTile',
        '    ├── reader/             # ReaderScreen, ReaderChrome, scrubber',
        '    ├── downloads/          # DownloadsScreen',
        '    └── settings/           # SettingsScreen, BackupRestoreScreen',
    ])

    story += h2('3.3  State Management')
    story += table(
        ['Provider type', 'Used for'],
        [
            ['StateProvider', 'Simple values: brightness, reading direction, scale, background, search query'],
            ['StreamProvider.family', 'Isar reactive streams — library list, chapter status, live manga object'],
            ['FutureProvider.family', 'Async loads — manga detail, chapter list, page URLs'],
            ['NotifierProvider', 'Complex mutation — library CRUD, category management'],
            ['AsyncNotifierProvider', 'Download manager — queue, pause/resume/cancel, progress tracking'],
        ],
        [CW*0.35, CW*0.65]
    )
    story += info(
        'ReaderState overrides == and hashCode so that Riverpod\'s equality check prevents '
        'redundant rebuilds when only unrelated fields change.'
    )

    story += h2('3.4  Database Schemas (Isar v3)')
    story += table(
        ['Schema', 'Key fields'],
        [
            ['MangaEntry', 'id, sourceId, sourceMangaId, title, coverUrl, author, genres, status, inLibrary, chapterCount, unreadCount, categories, lastUpdated'],
            ['ChapterEntry', 'id, mangaId, sourceChapterId, title, number, volume, language, uploadDate, isRead, isDownloaded'],
            ['DownloadEntry', 'id, mangaId, chapterId, mangaTitle, chapterNumber, status, progress, downloadedPages, totalPages'],
        ],
        [CW*0.24, CW*0.76]
    )

    story += h2('3.5  Theming')
    story.append(body('The AppColorsX extension on BuildContext resolves every colour at runtime:'))
    story += code_block([
        '// Use in any build method:',
        'final bg       = context.backgroundColor;       // dark: #0A0A0F   light: #F2F2F7',
        'final surface  = context.surfaceColor;           // dark: #141418   light: #FFFFFF',
        'final elevated = context.surfaceElevatedColor;   // dark: #1C1C24   light: #ECECF0',
        'final border   = context.borderColor;            // dark: 8% white  light: 8% black',
        'final textSec  = context.textSecondaryColor;     // 60% opacity',
    ])
    story.append(PageBreak())

    # ── Section 4: Source Plugin System ─────────────────────────────────
    story += h1('4', 'Source Plugin System')

    story += h2('4.1  MangaSource Interface')
    story += code_block([
        'abstract class MangaSource {',
        '  String get id;         // unique slug, e.g. "asurascans_en"',
        '  String get name;       // display name',
        '  String get baseUrl;    // root URL for web links',
        '  String get language;   // ISO 639-1 code',
        '  Map<String, String> get imageHeaders;',
        '  List<SourceFilter> getFilters();',
        '',
        '  Future<List<MangaSummary>> fetchPopular({int page = 1});',
        '  Future<List<MangaSummary>> fetchLatestUpdates({int page = 1});',
        '  Future<List<MangaSummary>> search(String query, ...);',
        '  Future<MangaDetail> fetchMangaDetail(String mangaId);',
        '  Future<List<ChapterInfo>> fetchChapterList(String mangaId);',
        '  Future<List<String>> fetchPageUrls(String chapterId);',
        '}',
    ])

    story += h2('4.2  Built-in Sources')
    story += h3('AsuraScans (asurascans_en)')
    story.append(bullets([
        f'API base: {code_inline("https://api.asurascans.com/api")}',
        f'Chapter ID format: {code_inline("seriesSlug::chapterUUID::chapterNumber")}',
        'fetchPageUrls: 3-strategy fallback (UUID path → chapter-number query → direct UUID endpoint)',
        'Sends Origin and Referer headers with every request.',
    ]))
    story += h3('AllManga (all_manga_en)')
    story.append(bullets([
        f'API base: {code_inline("https://api.allanime.day/api")} (GraphQL)',
        f'Chapter ID format: {code_inline("mangaId::chapterNumber")}',
        'Tries POST first (standard GraphQL), falls back to GET on DioException.',
    ]))

    story += h2('4.3  Adding a New Source')
    story.append(body('1. Create <i>lib/core/extensions/sources/my_source.dart</i> and implement MangaSource.'))
    story.append(body('2. Register your instance in <i>lib/core/providers/source_registry_provider.dart</i>.'))
    story += info(
        'Image headers must include Referer and Origin — many CDNs reject requests without them.',
        warn=True
    )
    story.append(PageBreak())

    # ── Section 5: Getting Started ───────────────────────────────────────
    story += h1('5', 'Getting Started')

    story += h2('5.1  Prerequisites')
    story += table(
        ['Requirement', 'Version'],
        [
            ['Flutter SDK', '≥ 3.3.0'],
            ['Dart SDK', '≥ 3.3.0'],
            ['Android SDK', 'API 21+ (Android 5.0 Lollipop)'],
            ['Java / JDK', '17 (required by Gradle 8)'],
        ],
        [CW*0.40, CW*0.60]
    )

    story += h2('5.2  Build & Run')
    story += h3('1. Clone and fetch dependencies')
    story += code_block([
        'git clone https://github.com/tsnyders/comic-center.git',
        'cd comic-center',
        'flutter pub get',
    ])
    story += h3('2. Generate code')
    story.append(body('Isar schemas and Riverpod annotations require code generation:'))
    story += code_block(['dart run build_runner build --delete-conflicting-outputs'])
    story += info('Never manually edit *.g.dart or *.freezed.dart files — they are regenerated on every build_runner run.', warn=True)
    story += h3('3. Run & Analyse')
    story += code_block([
        '# Debug on a connected device or emulator',
        'flutter run',
        '',
        '# Release build (optimised APK)',
        'flutter build apk --release',
        '',
        '# Static analysis (expect info-level warnings only)',
        'flutter analyze',
    ])
    story += info(
        'Critical: Do NOT change the applicationId (com.comiccenter.comic_center). '
        'Changing it forces users to uninstall and reinstall, losing all local data.',
        warn=True
    )
    story.append(PageBreak())

    # ── Section 6: Settings Reference ───────────────────────────────────
    story += h1('6', 'Settings Reference')
    story.append(body('All settings are persisted via Riverpod StateProviders backed by SharedPreferences. Changes take effect immediately.'))

    story += h2('Reader')
    story += table(
        ['Setting', 'Provider', 'Options', 'Default'],
        [
            ['Theme', 'brightnessProvider', 'dark / light', 'dark'],
            ['Reading Direction', 'readingDirectionProvider', 'ltr / rtl / vertical', 'ltr'],
            ['Page Scale', 'pageScaleModeProvider', 'fitWidth / fitHeight / original', 'fitWidth'],
            ['Background', 'readerBackgroundProvider', 'black / white / sepia', 'black'],
            ['Download Location', 'downloadLocationProvider', 'local / googleDrive', 'local'],
        ],
        [CW*0.26, CW*0.36, CW*0.26, CW*0.12]
    )

    story += h2('Update Check')
    story.append(body('Calls: https://api.github.com/repos/tsnyders/comic-center/releases/latest'))
    story.append(bullets([
        '<b>HTTP 200</b> — release exists. Displays version tag and changelog body. Download button if APK asset is attached.',
        '<b>HTTP 404</b> — no releases yet. Shows "No Updates Available".',
        '<b>Network error</b> — throws UpdateCheckException. Shows "Error" with specific message.',
    ]))
    story.append(PageBreak())

    # ── Section 7: Backup & Data ─────────────────────────────────────────
    story += h1('7', 'Backup & Data')

    story += h2('Backup File Format')
    story.append(body('Backups are stored as UTF-8 JSON at:'))
    story += code_block(['{appDocumentsDirectory}/backups/yomi_backup_YYYYMMDD_HHmmss.json'])
    story.append(body('Top-level structure:'))
    story += code_block([
        '{',
        '  "version": 1,',
        '  "exportedAt": "2026-06-07T10:30:00Z",',
        '  "categories": ["Favourites", "Reading"],',
        '  "manga": [',
        '    {',
        '      "title": "Solo Leveling",',
        '      "sourceId": "asurascans_en",',
        '      "sourceMangaId": "solo-leveling-abc123",',
        '      "categories": ["Favourites"],',
        '      "chapters": [',
        '        { "number": 1, "isRead": true, "isDownloaded": false }',
        '      ]',
        '    }',
        '  ]',
        '}',
    ])

    story += h2('Restore Behaviour')
    story.append(body('The restore operation performs an upsert merge:'))
    story.append(bullets([
        'Find the existing record by sourceId + sourceMangaId.',
        'If found, update metadata — <b>preserve</b> isRead and isDownloaded on chapters.',
        'If not found, insert as a new entry.',
        'Categories in the backup are merged with existing ones (no duplicates).',
    ]))
    story += info('A restore never deletes existing data. To start fresh, clear app data from Android Settings first.')

    story += h2('Pull-to-Refresh')
    story.append(body('Pulling down on the chapter list triggers refreshMangaChapters:'))
    story.append(bullets([
        'New chapters are inserted.',
        'Existing chapters have title, number, and uploadDate updated.',
        'isRead and isDownloaded are <b>never</b> modified.',
        'chapterCount and unreadCount on MangaEntry are recalculated.',
    ]))

    return story


# ── Page template callbacks ────────────────────────────────────────────────
def on_first_page(canvas, doc):
    canvas.saveState()
    draw_cover(canvas)
    canvas.restoreState()


def _draw_header_footer(canvas, page_num):
    canvas.saveState()
    canvas.setFont('Helvetica', 8)
    canvas.setFillColor(TEXT_MUTE)
    canvas.drawString(2.5*cm, H - 1.6*cm, 'Yomi — Technical Documentation')
    canvas.setStrokeColor(BORDER)
    canvas.setLineWidth(0.4)
    canvas.line(2.5*cm, H - 1.8*cm, W - 2.5*cm, H - 1.8*cm)
    canvas.line(2.5*cm, 1.5*cm, W - 2.5*cm, 1.5*cm)
    canvas.drawRightString(W - 2.5*cm, 1.0*cm, f'Page {page_num}')
    canvas.restoreState()


def on_later_pages(canvas, doc):
    _draw_header_footer(canvas, doc.page - 1)


def build_pdf(out_path):
    doc = SimpleDocTemplate(
        out_path,
        pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.2*cm, bottomMargin=2.2*cm,
        title='Yomi — Technical Documentation',
        author='Yomi Project',
    )
    story = build_story()
    doc.build(story, onFirstPage=on_first_page, onLaterPages=on_later_pages)
    print(f'PDF written to {out_path}')


# ─────────────────────────────────────────────────────────────────────────
#  DOCX via python-docx
# ─────────────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _shd(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def _para_shd(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    pPr.append(shd)


def build_docx(out_path):
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)

    # ── Helpers ────────────────────────────────────────────────────────────

    def dh1(num, title):
        p = doc.add_heading(level=1)
        p.clear()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(f'{num}  {title}')
        r.font.name = 'Calibri'; r.font.size = Pt(22)
        r.font.color.rgb = RGBColor(10, 10, 31); r.bold = True

    def dh2(text):
        p = doc.add_heading(level=2)
        p.clear()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(text)
        r.font.name = 'Calibri'; r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(10, 10, 31); r.bold = True

    def dh3(text):
        p = doc.add_heading(level=3)
        p.clear()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.font.name = 'Calibri'; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(10, 10, 31); r.bold = True

    def dbody(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        r = p.add_run(text)
        r.font.name = 'Calibri'; r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(42, 42, 62)

    def dbullet(items):
        for item in items:
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(item)
            r.font.name = 'Calibri'; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(42, 42, 62)

    def dcode(lines):
        for line in lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.left_indent  = Cm(0.5)
            _para_shd(p, '#0F0F1A')
            r = p.add_run(line if line else ' ')
            r.font.name = 'Courier New'; r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(200, 200, 232)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def dinfo(text, warn=False):
        fill = '#FFF3CD' if warn else '#EBF5FF'
        fg   = RGBColor(100, 60, 0) if warn else RGBColor(0, 60, 120)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Cm(0.5)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        _para_shd(p, fill)
        prefix = '⚠️  ' if warn else 'ℹ️  '
        r = p.add_run(prefix + text)
        r.font.name = 'Calibri'; r.font.size = Pt(10)
        r.font.color.rgb = fg

    def dtable(headers, rows, col_widths_cm):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = 'Table Grid'
        # header
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.width = Cm(col_widths_cm[i])
            _shd(cell, '#0A84FF')
            r = cell.paragraphs[0].add_run(h)
            r.font.name = 'Calibri'; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(255, 255, 255); r.bold = True
        # data
        for ri, row in enumerate(rows):
            bg = '#F5F5FC' if ri % 2 else '#FFFFFF'
            for i, val in enumerate(row):
                cell = t.rows[ri + 1].cells[i]
                cell.width = Cm(col_widths_cm[i])
                _shd(cell, bg)
                r = cell.paragraphs[0].add_run(str(val))
                r.font.name = 'Calibri'; r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(42, 42, 62)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Cover ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(72)
    r = p.add_run('Yomi')
    r.font.name = 'Calibri'; r.font.size = Pt(52); r.bold = True
    r.font.color.rgb = RGBColor(10, 10, 31)

    p2 = doc.add_paragraph()
    r2 = p2.add_run('Manga & manhwa reader for Android\nDeveloper and user reference')
    r2.font.name = 'Calibri'; r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(80, 80, 110)

    p3 = doc.add_paragraph()
    r3 = p3.add_run('Version 1.0.0  ·  Flutter 3.x  ·  Dart 3.3')
    r3.font.name = 'Calibri'; r3.font.size = Pt(11); r3.bold = True
    r3.font.color.rgb = RGBColor(10, 132, 255)

    p4 = doc.add_paragraph()
    p4.paragraph_format.space_before = Pt(12)
    r4 = p4.add_run('Platform: Android API 21+  |  UI: Cupertino  |  State: Riverpod v2  |  DB: Isar v3')
    r4.font.name = 'Calibri'; r4.font.size = Pt(9)
    r4.font.color.rgb = RGBColor(120, 120, 140)

    doc.add_page_break()

    # ── Section 1 ──────────────────────────────────────────────────────────
    dh1('1', 'Overview')
    dh2('1.1  What is Yomi?')
    dbody(
        'Yomi (読み — Japanese for "reading") is a feature-complete manga and manhwa '
        'reader for Android, built with Flutter’s Cupertino widget set for an iOS-native feel.'
    )
    dbody(
        'It manages a personal library, downloads chapters for offline reading, connects to pluggable '
        'content sources, and backs up reading history to local JSON or Google Drive.'
    )
    dh2('1.2  Core Design Principles')
    dbullet([
        'Immersive, minimal UI — dark-first with frosted glass and tap-to-reveal chrome.',
        'Reactive, no loading flickers — Isar streams + shimmer placeholders + Riverpod.',
        'Pluggable sources — MangaSource interface with zero build-time coupling.',
        'Adaptive light & dark — AppColorsX context extension resolves colours at runtime.',
    ])
    doc.add_page_break()

    # ── Section 2 ──────────────────────────────────────────────────────────
    dh1('2', 'Features')
    dh2('2.1  Library')
    dbullet([
        'Real-time search — frosted search bar filters titles as you type.',
        'Filter sheet — narrow by read status and genre.',
        'Categories — user-defined shelves; each manga can belong to multiple.',
        'Unread badge — shown on card corner for unread chapters.',
        'Shimmer skeleton — animated placeholder during database load.',
    ])
    dh2('2.2  Reader')
    dtable(
        ['Mode', 'Scroll', 'Use case'],
        [['Left → Right', 'Horizontal', 'Western / manhwa'],
         ['Right → Left', 'Horizontal reversed', 'Japanese manga'],
         ['Vertical', 'Vertical', 'Webtoons']],
        [4.0, 3.5, 7.0]
    )
    dbullet([
        'Tap to toggle chrome (toolbar + scrubber).',
        'Drag or tap the scrubber to jump to any page.',
        '2 pt progress line always visible at top of screen.',
        'Page-count pill after each turn, fades after 2 s.',
    ])
    dh2('2.3  Downloads')
    dbullet([
        'Queue — one at a time with live progress bar.',
        'Pause & Resume — resumes from where it left off.',
        'Cancel — partial files discarded.',
        'Retry — for failed chapters.',
    ])
    dtable(
        ['Status', 'Meaning'],
        [['Queued', 'Waiting in queue'],
         ['Downloading', 'In progress (spinner)'],
         ['Paused', 'Manually paused'],
         ['Failed', 'Failed — tap to retry'],
         ['Downloaded', 'Offline (green tick)']],
        [4.0, 10.5]
    )
    dh2('2.4  Browse & Sources')
    dtable(
        ['Source', 'Language', 'API'],
        [['AsuraScans', 'English', 'JSON REST — api.asurascans.com'],
         ['AllManga', 'English', 'GraphQL — api.allanime.day']],
        [4.0, 2.8, 7.7]
    )
    dh2('2.5  Backup & Restore')
    dbullet([
        'Export — timestamped JSON to app documents.',
        'Restore — upsert merge; preserves read/download status.',
    ])
    dinfo('Backup files are plain JSON and can be edited with any text editor.')
    dh2('2.6  Settings')
    dtable(
        ['Section', 'Option', 'Values'],
        [['Appearance', 'Theme', 'Dark / Light'],
         ['Reader', 'Direction', 'L→R / R→L / Vertical'],
         ['Reader', 'Page Scale', 'Fit Width / Fit Height / Original'],
         ['Reader', 'Background', 'Black / White / Sepia'],
         ['Downloads', 'Storage', 'Local / Google Drive'],
         ['Extensions', 'Check for Updates', 'GitHub release fetch'],
         ['Backup', 'Export / Restore', 'JSON to app documents']],
        [3.0, 3.8, 7.7]
    )
    doc.add_page_break()

    # ── Section 3 ──────────────────────────────────────────────────────────
    dh1('3', 'Architecture')
    dh2('3.1  Tech Stack')
    dtable(
        ['Layer', 'Technology', 'Version'],
        [['UI framework', 'Flutter — Cupertino', '≥ 3.3'],
         ['State', 'Riverpod', '^2.5'],
         ['Database', 'Isar', '^3.1'],
         ['Networking', 'Dio', '^5.4'],
         ['Images', 'extended_image', '^8.3'],
         ['Google Drive', 'googleapis + google_sign_in', '^12 / ^6.2'],
         ['Platform bridge', 'MethodChannel yomi/platform', '—']],
        [4.5, 6.5, 3.5]
    )
    dh2('3.2  Project Structure')
    dcode([
        'lib/',
        '├── core/',
        '│   ├── database/          # Isar schemas',
        '│   ├── extensions/sources/ # MangaSource plugins',
        '│   ├── providers/          # Riverpod providers',
        '│   ├── services/           # BackupService, UpdateService',
        '│   └── theme/              # AppColors, AppColorsX',
        '└── features/',
        '    ├── library/  browse/  title_detail/',
        '    ├── reader/  downloads/  settings/',
    ])
    dh2('3.3  State Management')
    dtable(
        ['Provider type', 'Used for'],
        [['StateProvider', 'brightness, direction, scale, background, search'],
         ['StreamProvider.family', 'Isar streams — library, chapter status, live manga'],
         ['FutureProvider.family', 'manga detail, chapters, page URLs'],
         ['NotifierProvider', 'Library CRUD, categories'],
         ['AsyncNotifierProvider', 'Download queue, pause/resume/cancel']],
        [4.5, 10.0]
    )
    dinfo('ReaderState overrides == and hashCode to prevent spurious rebuilds.')
    dh2('3.4  Database Schemas')
    dtable(
        ['Schema', 'Key fields'],
        [['MangaEntry', 'id, sourceId, sourceMangaId, title, coverUrl, author, genres, status, inLibrary, chapterCount, categories'],
         ['ChapterEntry', 'id, mangaId, sourceChapterId, title, number, uploadDate, isRead, isDownloaded'],
         ['DownloadEntry', 'id, mangaId, chapterId, status, progress, downloadedPages, totalPages']],
        [3.5, 11.0]
    )
    dh2('3.5  Theming')
    dcode([
        'context.backgroundColor       // dark: #0A0A0F   light: #F2F2F7',
        'context.surfaceColor           // dark: #141418   light: #FFFFFF',
        'context.surfaceElevatedColor   // dark: #1C1C24   light: #ECECF0',
        'context.borderColor            // dark: 8% white  light: 8% black',
    ])
    doc.add_page_break()

    # ── Section 4 ──────────────────────────────────────────────────────────
    dh1('4', 'Source Plugin System')
    dh2('4.1  MangaSource Interface')
    dcode([
        'abstract class MangaSource {',
        '  String get id;  String get name;  String get baseUrl;',
        '  String get language;',
        '  Map<String, String> get imageHeaders;',
        '  List<SourceFilter> getFilters();',
        '  Future<List<MangaSummary>> fetchPopular({int page = 1});',
        '  Future<List<MangaSummary>> fetchLatestUpdates({int page = 1});',
        '  Future<List<MangaSummary>> search(String query, ...);',
        '  Future<MangaDetail> fetchMangaDetail(String mangaId);',
        '  Future<List<ChapterInfo>> fetchChapterList(String mangaId);',
        '  Future<List<String>> fetchPageUrls(String chapterId);',
        '}',
    ])
    dh2('4.2  Built-in Sources')
    dh3('AsuraScans')
    dbullet([
        'API: https://api.asurascans.com/api',
        'Chapter ID: seriesSlug::chapterUUID::chapterNumber',
        '3-strategy fetchPageUrls fallback (UUID → chapter number → direct UUID)',
    ])
    dh3('AllManga')
    dbullet([
        'API: https://api.allanime.day/api (GraphQL)',
        'Chapter ID: mangaId::chapterNumber',
        'POST-first GraphQL with GET fallback on DioException.',
    ])
    dh2('4.3  Adding a New Source')
    dbody('1. Create lib/core/extensions/sources/my_source.dart, implement MangaSource.')
    dbody('2. Register in lib/core/providers/source_registry_provider.dart.')
    dinfo('Include Referer and Origin in imageHeaders — many CDNs reject requests without them.', warn=True)
    doc.add_page_break()

    # ── Section 5 ──────────────────────────────────────────────────────────
    dh1('5', 'Getting Started')
    dtable(
        ['Requirement', 'Version'],
        [['Flutter SDK', '≥ 3.3.0'], ['Dart SDK', '≥ 3.3.0'],
         ['Android SDK', 'API 21+'], ['Java / JDK', '17']],
        [6.0, 8.5]
    )
    dh2('Build & Run')
    dcode(['git clone https://github.com/tsnyders/comic-center.git',
           'cd comic-center && flutter pub get'])
    dcode(['dart run build_runner build --delete-conflicting-outputs'])
    dinfo('Never manually edit *.g.dart or *.freezed.dart files.', warn=True)
    dcode(['flutter run                    # debug',
           'flutter build apk --release    # release APK',
           'flutter analyze                # expect info-level only'])
    dinfo('Do NOT change applicationId (com.comiccenter.comic_center) — forces users to reinstall.', warn=True)
    doc.add_page_break()

    # ── Section 6 ──────────────────────────────────────────────────────────
    dh1('6', 'Settings Reference')
    dtable(
        ['Setting', 'Provider', 'Options', 'Default'],
        [['Theme', 'brightnessProvider', 'dark/light', 'dark'],
         ['Direction', 'readingDirectionProvider', 'ltr/rtl/vertical', 'ltr'],
         ['Page Scale', 'pageScaleModeProvider', 'fitWidth/fitHeight/original', 'fitWidth'],
         ['Background', 'readerBackgroundProvider', 'black/white/sepia', 'black'],
         ['Storage', 'downloadLocationProvider', 'local/googleDrive', 'local']],
        [3.5, 5.0, 4.5, 1.5]
    )
    dh2('Update Check')
    dbullet([
        'HTTP 200 — release found. Shows tag & changelog. Download button if APK attached.',
        'HTTP 404 — no releases yet. Shows "No Updates Available".',
        'Network error — throws UpdateCheckException. Shows error message.',
    ])
    doc.add_page_break()

    # ── Section 7 ──────────────────────────────────────────────────────────
    dh1('7', 'Backup & Data')
    dh2('Backup File Format')
    dcode(['{appDocumentsDirectory}/backups/yomi_backup_YYYYMMDD_HHmmss.json'])
    dcode(['{ "version": 1, "exportedAt": "...", "categories": [...],',
           '  "manga": [{ "title": "...", "sourceId": "...",',
           '    "chapters": [{ "number": 1, "isRead": true }] }] }'])
    dh2('Restore Behaviour')
    dbullet([
        'Find existing record by sourceId + sourceMangaId.',
        'If found, update metadata — preserve isRead and isDownloaded.',
        'If not found, insert as new entry.',
        'Categories merged (no duplicates).',
    ])
    dinfo('A restore never deletes data. Clear app data from Android Settings to start fresh.')
    dh2('Pull-to-Refresh')
    dbullet([
        'New chapters are inserted.',
        'Existing chapters: title, number, uploadDate updated; isRead/isDownloaded unchanged.',
        'chapterCount and unreadCount recalculated on MangaEntry.',
    ])

    doc.save(out_path)
    print(f'DOCX written to {out_path}')


# ─────────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    build_pdf(os.path.join(OUT_DIR, 'yomi_documentation.pdf'))
    build_docx(os.path.join(OUT_DIR, 'yomi_documentation.docx'))
    print('Done.')
